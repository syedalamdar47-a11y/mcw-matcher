"""
sop_docx.py — reusable python-docx toolkit for building SOPs in the MCW house style.

Drop this file next to your build script and import from it. It encapsulates every
raw-OXML trick the house style needs (real PAGE/TOC fields, cell shading and borders,
clickable hyperlinks, [[key|display]] link markers) so build scripts stay readable.

Requires: python-docx  (pip install python-docx)
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------
# House-style constants — do not change these without updating the spec.
# --------------------------------------------------------------------------
NAVY = "1F3864"          # headings, table header fills, label runs
GREY = "595959"          # page header/footer, title-page supporting text
BORDER_GREY = "BFBFBF"   # standard table cell borders
IMAGE_BORDER = "9E9E9E"  # screenshot frame border
LINK_BLUE = "0563C1"     # hyperlink text

WARN_FILL = "FFF2CC"     # "Watch out" callout
TIP_FILL = "DEEBF7"      # "Tip" callout
OK_FILL = "E2EFDA"       # "Success / acceptance note" callout

CALLOUTS = {
    "warn": ("⚠  Watch out:", WARN_FILL),
    "tip": ("\U0001F4A1  Tip:", TIP_FILL),
    "ok": ("✅  Tip:", OK_FILL),
}

BASE_FONT = "Arial"
BASE_SIZE_PT = 11
HEADING_SIZES = {1: 18, 2: 14, 3: 12}
IMAGE_WIDTH_IN = 6.2
CALLOUT_WIDTH_IN = 6.4

MARKER_RE = re.compile(r"\[\[(\w+)\|([^\]]+)\]\]")


# --------------------------------------------------------------------------
# Low-level OXML helpers
# --------------------------------------------------------------------------
def _force_font(style_element, font_name: str) -> None:
    """Set w:rFonts on all four scripts so Word does not substitute the font."""
    rpr = style_element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def shade_cell(cell, fill_hex: str) -> None:
    """w:shd requires w:val; omitting it produces OXML Word may reject."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = BORDER_GREY, sz: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def add_page_field(paragraph, size_pt: int = 9, color: str = GREY):
    """Real Word PAGE field: fldChar begin + instrText + fldChar end."""
    run = paragraph.add_run()
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    return run


def add_toc_field(paragraph, placeholder: str) -> None:
    """Real Word TOC field: begin + instrText + separate + placeholder + end."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, sep, text, end):
        run._r.append(node)


def add_hyperlink(paragraph, text: str, url: str,
                  font_name: str = BASE_FONT, size_pt: int = BASE_SIZE_PT):
    """Blue underlined external hyperlink run."""
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    r_pr.append(rfonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))  # OXML sizes are half-points
    r_pr.append(sz)

    run.append(r_pr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


# --------------------------------------------------------------------------
# SopDoc
# --------------------------------------------------------------------------
class SopDoc:
    """Builds one SOP document in the house style.

    links: the parsed links.json catalog (dict of key -> {name,type,access,url,...}).
    header_title: e.g. "SOP-003 — Client Intake"; appears in the page header.
    """

    def __init__(self, links: dict, header_title: str):
        self.links = {k: v for k, v in links.items() if not k.startswith("_")}
        self.header_title = header_title
        self.doc = Document()
        self.link_keys_used: list[str] = []  # first-use order, drives Tools table
        self._setup_styles()
        self._setup_section()

    # ---- setup ----
    def _setup_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = BASE_FONT
        normal.font.size = Pt(BASE_SIZE_PT)
        _force_font(normal.element, BASE_FONT)
        for level, size in HEADING_SIZES.items():
            s = self.doc.styles[f"Heading {level}"]
            s.font.name = BASE_FONT
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = RGBColor.from_string(NAVY)
            _force_font(s.element, BASE_FONT)

    def _setup_section(self) -> None:
        section = self.doc.sections[0]
        for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, side, Inches(1))

        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(self.header_title)
        hr.font.name = BASE_FONT
        hr.font.size = Pt(9)
        hr.font.color.rgb = RGBColor.from_string(GREY)

        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("Page ")
        fr.font.size = Pt(9)
        fr.font.color.rgb = RGBColor.from_string(GREY)
        add_page_field(fp)

    # ---- primitives ----
    def heading(self, text: str, level: int = 1):
        return self.doc.add_heading(text, level=level)

    def para(self, text: str = "", style: str | None = None):
        return self.doc.add_paragraph(text, style=style)

    def spacer(self, count: int = 1) -> None:
        for _ in range(count):
            self.doc.add_paragraph()

    def page_break(self) -> None:
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def resolve(self, key: str) -> dict:
        if key not in self.links:
            raise KeyError(
                f"Link marker [[{key}|...]] does not match any key in links.json. "
                f"Known keys: {', '.join(sorted(self.links))}"
            )
        if key not in self.link_keys_used:
            self.link_keys_used.append(key)
        return self.links[key]

    def rich(self, paragraph, text: str, size_pt: int = BASE_SIZE_PT) -> None:
        """Write text into a paragraph, converting [[key|display]] into hyperlinks.

        Raises KeyError on an unresolved marker rather than leaking the raw
        marker into the finished document.
        """
        last = 0
        for m in MARKER_RE.finditer(text):
            if m.start() > last:
                paragraph.add_run(text[last:m.start()])
            entry = self.resolve(m.group(1))
            add_hyperlink(paragraph, m.group(2), entry["url"], size_pt=size_pt)
            last = m.end()
        if last < len(text):
            paragraph.add_run(text[last:])
        leftover = re.search(r"\[\[[^\]]*\]\]", "".join(r.text for r in paragraph.runs))
        if leftover:
            raise ValueError(f"Malformed link marker survived parsing: {leftover.group(0)}")

    def rich_para(self, text: str):
        p = self.doc.add_paragraph()
        self.rich(p, text)
        return p

    def labeled(self, label: str, text: str):
        """Bold navy label followed by body text with link-marker support."""
        p = self.doc.add_paragraph()
        r = p.add_run(label + " ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        self.rich(p, text)
        return p

    def bullets(self, items) -> None:
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            self.rich(p, item)

    # ---- blocks ----
    def screenshot(self, image_path, width_in: float = IMAGE_WIDTH_IN) -> None:
        """Embed an annotated screenshot inside a bordered 1x1 table."""
        image_path = Path(image_path)
        if not image_path.exists():
            raise FileNotFoundError(f"Screenshot not found: {image_path}")
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_borders(cell, color=IMAGE_BORDER, sz="8")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(width_in))

    def callout(self, kind: str, text: str) -> None:
        """kind: 'warn' | 'tip' | 'ok'."""
        if kind not in CALLOUTS:
            raise ValueError(f"Unknown callout kind {kind!r}; use one of {list(CALLOUTS)}")
        label, fill = CALLOUTS[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(CALLOUT_WIDTH_IN)
        shade_cell(cell, fill)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        r1 = p.add_run(label + "  ")   # two spaces separate label from body
        r1.bold = True                 # label keeps the default 11 pt
        r2 = p.add_run(text)
        r2.font.size = Pt(10)

    def step(self, number: int, title: str, image_path, do: str, see: str,
             callout: tuple[str, str] | None = None) -> None:
        """One complete procedure step in house order."""
        self.heading(f"Step {number}: {title}", level=3)
        if image_path is not None:
            self.screenshot(image_path)
        self.labeled("What to do:", do)
        self.labeled("What you should see:", see)
        if callout:
            self.callout(callout[0], callout[1])

    # ---- fixed sections ----
    def title_page(self, sop_number: str, short_title: str, subtitle: str,
                   version: str, date_created: str) -> None:
        self.spacer(6)
        for text, bold, size, color, italic in (
            ("STANDARD OPERATING PROCEDURE", True, 14, GREY, False),
            (sop_number, True, 28, NAVY, False),
            (short_title, True, 24, NAVY, False),
        ):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor.from_string(color)

        self.spacer(1)  # single spacer between title and subtitle
        sp = self.doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(subtitle)
        sr.italic = True
        sr.font.size = Pt(12)
        sr.font.color.rgb = RGBColor.from_string(GREY)

        self.spacer(6)
        meta = self.doc.add_table(rows=4, cols=2)
        meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows = [("Version", version), ("Date Created", date_created),
                ("Prepared by", ""), ("Process owner", "")]
        for i, (k, v) in enumerate(rows):
            c0, c1 = meta.cell(i, 0), meta.cell(i, 1)
            c0.width, c1.width = Inches(2), Inches(3)
            shade_cell(c0, NAVY)
            set_cell_borders(c0)
            set_cell_borders(c1)
            r0 = c0.paragraphs[0].add_run(k)
            r0.bold = True
            r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c1.paragraphs[0].add_run(v)
        self.page_break()

    def table_of_contents(self) -> None:
        self.heading("Table of Contents", level=1)
        add_toc_field(
            self.doc.add_paragraph(),
            "Right-click and select 'Update Field' to populate the Table of Contents.",
        )
        self.page_break()

    def tools_table(self, keys) -> None:
        """4-column Tool | Type | Access Needed | Open table, one row per key."""
        keys = list(keys)
        table = self.doc.add_table(rows=len(keys) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Tool", "Type", "Access Needed", "Open"]):
            c = table.cell(0, i)
            shade_cell(c, NAVY)
            set_cell_borders(c)
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
        for row_idx, key in enumerate(keys, start=1):
            entry = self.resolve(key)
            for col_idx, val in enumerate((entry["name"], entry["type"], entry["access"])):
                c = table.cell(row_idx, col_idx)
                set_cell_borders(c)
                r = c.paragraphs[0].add_run(val)
                r.font.size = Pt(10)
            c = table.cell(row_idx, 3)
            set_cell_borders(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_hyperlink(p, "↗ Open", entry["url"], size_pt=10)
        widths = [Inches(2.8), Inches(1.2), Inches(1.1), Inches(0.8)]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]
        self.spacer(1)  # spacer paragraph after the table

    def glossary(self, entries) -> None:
        """entries: iterable of (term, link_key_or_None, definition)."""
        for term, link_key, definition in entries:
            p = self.doc.add_paragraph()
            if link_key:
                entry = self.resolve(link_key)
                add_hyperlink(p, term, entry["url"])
                p.add_run(" — ")
            else:
                r = p.add_run(term + " — ")
                r.bold = True
            p.add_run(definition)

    def troubleshooting(self, pairs) -> None:
        """pairs: iterable of (question, answer)."""
        for question, answer in pairs:
            pq = self.doc.add_paragraph()
            rq = pq.add_run("Q: " + question)
            rq.bold = True
            self.doc.add_paragraph().add_run("A: " + answer)
            self.spacer(1)

    def revision_history(self, rows) -> None:
        """rows: iterable of (version, date, author, changes)."""
        rows = list(rows)
        table = self.doc.add_table(rows=len(rows) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Version", "Date", "Author", "Changes Made"]):
            c = table.cell(0, i)
            shade_cell(c, NAVY)
            set_cell_borders(c)
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row_idx, values in enumerate(rows, start=1):
            for i, v in enumerate(values):
                c = table.cell(row_idx, i)
                set_cell_borders(c)
                c.paragraphs[0].add_run(v)
        widths = [Inches(0.9), Inches(1.5), Inches(1.4), Inches(2.6)]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]

    # ---- output ----
    def save(self, path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        return path


# --------------------------------------------------------------------------
# QA helpers — run these before delivering (no Word installation needed).
# --------------------------------------------------------------------------
def verify_docx(path) -> dict:
    """Structural self-check of a built SOP. Returns a report dict; raises on hard failures."""
    import zipfile

    path = Path(path)
    report = {"path": str(path), "size_bytes": path.stat().st_size}
    with zipfile.ZipFile(path) as z:
        bad = z.testzip()
        if bad is not None:
            raise ValueError(f"Corrupt entry in docx: {bad}")
        names = z.namelist()
        document_xml = z.read("word/document.xml").decode("utf-8", "replace")
        footers = [n for n in names if n.startswith("word/footer")]
        footer_xml = "".join(z.read(n).decode("utf-8", "replace") for n in footers)
    report["has_toc_field"] = 'TOC \\o "1-3"' in document_xml
    report["has_page_field"] = "PAGE" in footer_xml
    report["images"] = len([n for n in names if n.startswith("word/media/")])
    report["hyperlinks"] = document_xml.count("<w:hyperlink")
    report["unresolved_markers"] = len(re.findall(r"\[\[[^\]]*\]\]", document_xml))
    Document(str(path))  # re-open round trip
    for key in ("has_toc_field", "has_page_field"):
        if not report[key]:
            raise ValueError(f"QA failed: {key} is False for {path}")
    if report["unresolved_markers"]:
        raise ValueError(f"QA failed: {report['unresolved_markers']} unresolved [[...]] markers in {path}")
    return report
